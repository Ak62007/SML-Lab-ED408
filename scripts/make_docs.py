"""
Generates the per-question ED408 lab documentation (.docx), following the
exact "Problem + Algorithm / Program / Output + Analysis" template format
supplied for this lab (see docs/reference template). Matches its style:
A4 page, Times New Roman 12pt, 1" margins. The "Program" section shows only
curated key code sections (function defs, core algorithm, key parameters)
per the template's instruction that "the complete executable program can
be maintained digitally through Github" -- full scripts live in scripts/.
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
METRICS = ROOT / "results" / "metrics"
FIGURES = ROOT / "results" / "figures"
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True, parents=True)

GITHUB_URL = "https://github.com/Ak62007/SML-Lab-ED408"

DATASET_NOTE = (
    "BSDS500 (Berkeley Segmentation Data Set) is an image-segmentation "
    "dataset, not tabular data, so a tabular dataset was derived from it "
    "(data/bsds_features.csv, built by scripts/build_dataset.py): 24,000 "
    "pixels sampled from the 200 BSDS500 training images, 10 features "
    "(R, G, B, gray, grad_mag, grad_dir, laplacian, local_std, x_norm, "
    "y_norm) capturing colour, gradient, texture and position, and label "
    "is_edge (1 if a majority of 6 human annotators marked the pixel as a "
    "segment boundary, else 0). This same dataset is used, unchanged, for "
    "every question in this lab."
)

BODY_SIZE = 12
CODE_SIZE = 9.5
TABLE_SIZE = 10


def load_json(name):
    with open(METRICS / f"{name}.json") as f:
        return json.load(f)


# ---------------------------------------------------------------- doc setup
def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(BODY_SIZE)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for section in doc.sections:
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    return doc


def tight(p, before=0, after=4, line=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    return p


def para(doc, text="", bold=False, italic=False, size=BODY_SIZE, after=4, align=None):
    p = doc.add_paragraph()
    tight(p, after=after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
    return p


def label_line(doc, label, value, after=4):
    p = doc.add_paragraph()
    tight(p, after=after)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(BODY_SIZE)
    r1.font.name = "Times New Roman"
    r2 = p.add_run(str(value))
    r2.font.size = Pt(BODY_SIZE)
    r2.font.name = "Times New Roman"
    return p


def section_heading(doc, n, text):
    p = doc.add_paragraph()
    tight(p, before=10, after=4)
    r = p.add_run(f"{n}. {text}")
    r.bold = True
    r.font.size = Pt(BODY_SIZE + 1)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x5C)
    return p


def sub_heading(doc, text):
    p = doc.add_paragraph()
    tight(p, before=6, after=3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(BODY_SIZE)
    r.font.name = "Times New Roman"
    return p


def numbered_list(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph()
        tight(p, after=2)
        r = p.add_run(f"{i}. {item}")
        r.font.size = Pt(BODY_SIZE)
        r.font.name = "Times New Roman"


def add_code(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    shd = cell._tc.get_or_add_tcPr()
    shade = shd.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "F2F2F2"})
    shd.append(shade)
    cell.paragraphs[0].text = ""
    for i, line in enumerate(code_text.strip("\n").splitlines()):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        tight(p, before=0, after=0, line=1.0)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(CODE_SIZE)
    return tbl


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for p in hdr[i].paragraphs:
            tight(p, after=0)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(TABLE_SIZE)
                r.font.name = "Times New Roman"
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                tight(p, after=0)
                for r in p.runs:
                    r.font.size = Pt(TABLE_SIZE)
                    r.font.name = "Times New Roman"
    return table


def add_image(doc, path, width_in=4.6, caption=None):
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(doc.paragraphs[-1], after=2)
    if caption:
        p = para(doc, caption, italic=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    return doc.paragraphs[-1]


# ---------------------------------------------------------------- builder
def build_doc(cfg):
    doc = new_doc()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(title_p, after=2)
    r = title_p.add_run(cfg["title"])
    r.bold = True
    r.font.size = Pt(15)
    r.font.name = "Times New Roman"

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(sub_p, after=10)
    r2 = sub_p.add_run("ED408 — Statistical / Supervised Machine Learning Lab")
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.name = "Times New Roman"

    # ============ 1. Problem + Algorithm ============
    section_heading(doc, 1, "Problem + Algorithm")
    label_line(doc, "Experiment No.", cfg["exp_no"])
    label_line(doc, "Title", cfg["title"])
    label_line(doc, "Aim", cfg["aim"])

    sub_heading(doc, "Problem Statement:")
    para(doc, cfg["problem_statement"], after=6)

    sub_heading(doc, "Algorithm:")
    numbered_list(doc, cfg["algorithm"])

    # ============ 2. Program ============
    section_heading(doc, 2, "Program")
    para(doc, "Key code section (function definitions, core algorithm, "
              "and key parameter settings):", after=3)
    add_code(doc, cfg["code"])
    p = para(doc, after=8)
    r1 = p.add_run("Full executable program: ")
    r1.font.size = Pt(BODY_SIZE); r1.font.name = "Times New Roman"
    r2 = p.add_run(f"{cfg['script_name']}")
    r2.italic = True; r2.font.size = Pt(BODY_SIZE); r2.font.name = "Times New Roman"
    r3 = p.add_run(" — maintained digitally on GitHub: ")
    r3.font.size = Pt(BODY_SIZE); r3.font.name = "Times New Roman"
    r4 = p.add_run(GITHUB_URL)
    r4.italic = True; r4.font.size = Pt(BODY_SIZE); r4.font.name = "Times New Roman"

    # ============ 3. Output + Analysis ============
    section_heading(doc, 3, "Output + Analysis")
    sub_heading(doc, "Parameters:")
    for label, value in cfg["parameters"]:
        label_line(doc, label, value, after=2)

    sub_heading(doc, "Program Execution:")
    for label, value in cfg["results"]:
        label_line(doc, label, value, after=2)
    label_line(doc, "Execution Status", "Successfully Executed", after=6)

    sub_heading(doc, "Result Graph:")
    para(doc, f"Figure: {cfg['graph_caption']}", bold=True, after=2)
    label_line(doc, "X-axis", cfg["x_axis"], after=1)
    label_line(doc, "Y-axis", cfg["y_axis"], after=4)
    add_image(doc, FIGURES / cfg["graph_file"], width_in=cfg.get("graph_width", 4.3))
    para(doc, cfg["graph_description"], after=8)

    sub_heading(doc, "Parameter Modification:")
    para(doc, cfg["param_mod_intro"], after=4)
    add_table(doc, cfg["param_mod_headers"], cfg["param_mod_rows"])
    para(doc, cfg["param_mod_conclusion"], after=4)

    out = DOCS / cfg["filename"]
    doc.save(out)
    print("wrote", out)


# =========================================================================
def q1_cfg():
    m = load_json("q1_knn")
    return dict(
        exp_no="01", title="k-Nearest Neighbours Classifier (From Scratch)",
        filename="K-Nearest Neighbours ED408.docx", script_name="scripts/q1_knn_scratch.py",
        aim="Implement the k-NN classification algorithm from scratch using only "
            "NumPy, and evaluate it on a real dataset.",
        problem_statement="Given a pixel's feature vector (colour, gradient, texture, "
            "position), classify it as an edge pixel or a non-edge pixel by finding its "
            "k nearest neighbours (Euclidean distance) in the standardized training "
            "feature space and taking a majority vote of their labels. " + DATASET_NOTE,
        algorithm=[
            "Load and standardize the feature dataset (zero mean, unit variance).",
            "Split data into train (80%) and test (20%) sets, stratified by class.",
            "Store all training feature vectors and their labels (no explicit training step).",
            "For each test point, compute the Euclidean distance to every training point.",
            "Select the k training points with the smallest distance.",
            "Predict the majority class label among those k neighbours.",
            "Repeat for every test point and compute test accuracy.",
            "Repeat the whole process for several values of k and pick the best k.",
        ],
        code='''class KNNFromScratch:
    def __init__(self, k=5):
        self.k = k

    def fit(self, X, y):
        self.X_train, self.y_train = X, y
        return self

    def predict(self, X, k=None):
        k = k or self.k
        train_sq = (self.X_train ** 2).sum(axis=1)
        batch_sq = (X ** 2).sum(axis=1, keepdims=True)
        # ||a-b||^2 = ||a||^2 + ||b||^2 - 2 a.b  (avoids O(n*m*d) memory)
        dists = batch_sq + train_sq[None, :] - 2 * X @ self.X_train.T
        nn_idx = np.argpartition(dists, kth=k - 1, axis=1)[:, :k]
        votes = self.y_train[nn_idx].sum(axis=1)
        return (votes > k / 2).astype(int)

# Key parameter setting: sweep k to find the best value
k_values = [1, 3, 5, 7, 9, 11, 15, 21, 31]
model = KNNFromScratch().fit(X_train, y_train)
accuracies = [accuracy_score(y_test, model.predict(X_test, k=k))
              for k in k_values]
best_k = k_values[int(np.argmax(accuracies))]''',
        parameters=[
            ("Feature Set", "10 features: R, G, B, gray, grad_mag, grad_dir, laplacian, local_std, x_norm, y_norm"),
            ("Train / Test Split", "19,200 / 4,800 samples (80/20, stratified)"),
            ("Distance Metric", "Euclidean"),
            ("k values tested", ", ".join(str(k) for k in m["k_values"])),
        ],
        results=[
            ("Best k", m["best_k"]),
            ("Test Accuracy", f"{m['test_accuracy']*100:.2f}%"),
            ("Confusion Matrix", str(m["confusion_matrix"])),
        ],
        graph_file="q1_knn_k_selection.png",
        graph_caption="Test accuracy vs. k for the from-scratch k-NN classifier.",
        x_axis="k (number of neighbours)", y_axis="Test accuracy",
        graph_description="The graph demonstrates that accuracy rises sharply from k=1 "
            "and plateaus around k=15-31, showing the classic bias-variance trade-off of k-NN.",
        param_mod_intro="Effect of varying k (the sole hyperparameter of k-NN) on test accuracy:",
        param_mod_headers=["k", "Test Accuracy"],
        param_mod_rows=[[k, f"{a:.4f}"] for k, a in zip(m["k_values"], m["accuracies"])],
        param_mod_conclusion="Small k (k=1) overfits to noisy individual pixels (lowest accuracy); "
            "accuracy improves and stabilizes as k grows, with the best result at "
            f"k={m['best_k']} ({m['test_accuracy']*100:.2f}%).",
    )


def q2_cfg():
    m = load_json("q2_decision_tree")
    top3 = sorted(m["feature_importances"].items(), key=lambda t: -t[1])[:3]
    sweep = m["depth_sweep"]
    return dict(
        exp_no="02", title="Decision Tree Classifier (scikit-learn) with Visualization",
        filename="Decision Tree ED408.docx", script_name="scripts/q2_decision_tree.py",
        aim="Build a Decision Tree classifier using scikit-learn and visualize the "
            "learned tree structure.",
        problem_statement="Learn a set of interpretable if-then rules over pixel colour, "
            "gradient, texture and position features that split the data into edge and "
            "non-edge pixels, and visualize the resulting decision tree. " + DATASET_NOTE,
        algorithm=[
            "Load the feature dataset (unscaled; trees are scale-invariant).",
            "Split data into train (80%) and test (20%) sets.",
            "At the root, evaluate every feature/threshold split and choose the one "
            "that most reduces Gini impurity.",
            "Partition the data into left and right child nodes at that split.",
            "Recurse on each child node until max_depth or min_samples_leaf is reached.",
            "Assign each leaf the majority class of the training samples that reach it.",
            "Predict new samples by traversing the tree from root to a leaf.",
            "Visualize the trained tree structure and rank feature importances.",
        ],
        code='''from sklearn.tree import DecisionTreeClassifier, plot_tree

clf = DecisionTreeClassifier(
    max_depth=6, min_samples_leaf=20, random_state=42)
clf.fit(X_train, y_train)
preds = clf.predict(X_test)
acc = accuracy_score(y_test, preds)

# feature importances (Gini-based)
importances = sorted(zip(FEATURE_COLS, clf.feature_importances_),
                      key=lambda t: -t[1])

# tree visualization
plot_tree(clf, feature_names=FEATURE_COLS,
          class_names=["non-edge", "edge"],
          filled=True, rounded=True, fontsize=10)''',
        parameters=[
            ("Feature Set", "10 features (unscaled)"),
            ("Train / Test Split", "19,200 / 4,800 samples (80/20)"),
            ("Splitting Criterion", "Gini impurity"),
            ("max_depth", 6),
            ("min_samples_leaf", 20),
        ],
        results=[
            ("Test Accuracy (depth=6)", f"{m['depth6_test_accuracy']*100:.2f}%"),
            ("Test Accuracy (depth=3, readable tree)", f"{m['depth3_test_accuracy']*100:.2f}%"),
            ("Top Feature", f"{top3[0][0]} (importance={top3[0][1]:.3f})"),
        ],
        graph_file="q2_tree_shallow.png",
        graph_caption="Visualized Decision Tree (max_depth=3, readable).",
        x_axis="Feature splits (root to leaves)", y_axis="Tree depth",
        graph_description="Each box shows the split condition, Gini impurity, sample count "
            "and majority class; local_std (local texture energy) dominates the top splits.",
        graph_width=5.4,
        param_mod_intro="Effect of varying max_depth (the key regularization hyperparameter) "
            "on test accuracy:",
        param_mod_headers=["max_depth", "Test Accuracy"],
        param_mod_rows=[[d, f"{a:.4f}"] for d, a in zip(sweep["depths"], sweep["accuracies"])],
        param_mod_conclusion="Accuracy improves as depth increases up to depth=6-8, then "
            "declines as the tree overfits (depth=15 and unbounded depth score lower than "
            "depth=6), confirming max_depth is an effective regularizer.",
    )


def q3_cfg():
    m = load_json("q3_naive_bayes")
    sweep = m["smoothing_sweep"]
    return dict(
        exp_no="03", title="Naive Bayes Classifier (Spam/Ham-style Binary Classification)",
        filename="Naive Bayes ED408.docx", script_name="scripts/q3_naive_bayes.py",
        aim="Apply a Naive Bayes classifier to a spam-vs-ham-style binary "
            "classification problem.",
        problem_statement="Classify each pixel as 'spam' (edge) or 'ham' (non-edge) "
            "using GaussianNB, under the naive assumption that the 10 colour/gradient/"
            "texture/position features are conditionally independent given the class. "
            "Per lab instructions the dataset used is not an email corpus but the same "
            "derived BSDS500 pixel dataset used throughout this lab; edge/non-edge here "
            "plays the role spam/ham would in a text classifier. " + DATASET_NOTE,
        algorithm=[
            "Load the feature dataset.",
            "Split data into train (80%) and test (20%) sets.",
            "For each class, estimate the mean and variance of every feature "
            "(fit a Gaussian per feature per class).",
            "Estimate class priors P(spam), P(ham) from training label frequencies.",
            "For a new sample, compute the likelihood of each feature value under "
            "each class's Gaussian.",
            "Multiply the per-feature likelihoods together with the class prior "
            "(naive independence assumption) to get an unnormalized posterior.",
            "Predict the class with the higher posterior probability.",
            "Evaluate accuracy, confusion matrix and ROC-AUC on the test set.",
        ],
        code='''from sklearn.naive_bayes import GaussianNB

clf = GaussianNB(var_smoothing=1e-9)
clf.fit(X_train, y_train)

preds = clf.predict(X_test)
probs = clf.predict_proba(X_test)[:, 1]

acc = accuracy_score(y_test, preds)
class_priors = clf.class_prior_   # P(ham), P(spam)''',
        parameters=[
            ("Feature Set", "10 features (unscaled)"),
            ("Train / Test Split", "19,200 / 4,800 samples (80/20)"),
            ("Likelihood Model", "Gaussian, per feature per class"),
            ("var_smoothing", "1e-9 (default)"),
        ],
        results=[
            ("Test Accuracy", f"{m['test_accuracy']*100:.2f}%"),
            ("Class Priors (ham, spam)", [round(v, 4) for v in m["class_priors"]]),
            ("Confusion Matrix", str(m["confusion_matrix"])),
        ],
        graph_file="q3_nb_roc.png",
        graph_caption="ROC curve for the Naive Bayes (GaussianNB) classifier.",
        x_axis="False Positive Rate", y_axis="True Positive Rate",
        graph_description="The curve sits well above the diagonal (random) line, showing "
            "the model ranks edge pixels above non-edge pixels substantially better than chance.",
        graph_width=3.6,
        param_mod_intro="Effect of varying var_smoothing (Laplace-style variance smoothing) "
            "on test accuracy:",
        param_mod_headers=["var_smoothing", "Test Accuracy"],
        param_mod_rows=[[v, f"{a:.4f}"] for v, a in zip(sweep["var_smoothing"], sweep["accuracies"])],
        param_mod_conclusion="Accuracy is stable for var_smoothing below 1e-5, then degrades "
            "as smoothing grows large (var_smoothing=1) because the per-class Gaussians are "
            "forced to look nearly identical, destroying discriminative power.",
    )


def q4_cfg():
    m = load_json("q4_logistic_regression")
    top3 = sorted(m["weights"].items(), key=lambda t: -abs(t[1]))[:3]
    sweep = m["lr_sweep"]
    return dict(
        exp_no="04", title="Logistic Regression (From Scratch, Gradient Descent)",
        filename="Logistic Regression ED408.docx", script_name="scripts/q4_logistic_regression.py",
        aim="Implement Logistic Regression from scratch using batch gradient descent "
            "to minimize the binary cross-entropy loss.",
        problem_statement="Learn a linear decision boundary (in standardized feature "
            "space) that separates edge pixels from non-edge pixels, by fitting weights "
            "w and bias b that minimize binary cross-entropy loss via gradient descent. "
            + DATASET_NOTE,
        algorithm=[
            "Load and standardize the feature dataset.",
            "Split data into train (80%) and test (20%) sets.",
            "Initialize weight vector w and bias b to zero.",
            "Compute predicted probabilities p = sigmoid(X.w + b) for the whole "
            "training set (batch gradient descent).",
            "Compute the gradient of the binary cross-entropy loss w.r.t. w and b.",
            "Update w, b in the direction that reduces the loss: w := w - lr * grad.",
            "Repeat steps 4-6 for n_iters iterations, tracking the loss.",
            "Threshold final probabilities at 0.5 to make class predictions.",
        ],
        code='''class LogisticRegressionScratch:
    def __init__(self, lr=0.5, n_iters=3000):
        self.lr, self.n_iters = lr, n_iters

    @staticmethod
    def _sigmoid(z):
        return 1.0 / (1.0 + np.exp(-z))

    def fit(self, X, y):
        n, d = X.shape
        Xb = np.hstack([np.ones((n, 1)), X])   # bias term
        self.w = np.zeros(d + 1)
        for it in range(self.n_iters):
            p = self._sigmoid(Xb @ self.w)
            grad = Xb.T @ (p - y) / n           # BCE gradient
            self.w -= self.lr * grad            # gradient descent step
        return self

    def predict(self, X, threshold=0.5):
        Xb = np.hstack([np.ones((len(X), 1)), X])
        return (self._sigmoid(Xb @ self.w) >= threshold).astype(int)

model = LogisticRegressionScratch(lr=0.5, n_iters=3000).fit(X_train, y_train)''',
        parameters=[
            ("Feature Set", "10 standardized features"),
            ("Train / Test Split", "19,200 / 4,800 samples (80/20)"),
            ("Learning Rate", 0.5),
            ("Iterations", 3000),
            ("Loss Function", "Binary Cross-Entropy"),
        ],
        results=[
            ("Test Accuracy", f"{m['test_accuracy']*100:.2f}%"),
            ("Final Training Loss", f"{m['final_loss']:.4f}"),
            ("Top Weight", f"{top3[0][0]} = {top3[0][1]:+.3f}"),
        ],
        graph_file="q4_logreg_loss_curve.png",
        graph_caption="Convergence of Logistic Regression via batch gradient descent.",
        x_axis="Iteration", y_axis="Binary cross-entropy loss",
        graph_description="The graph demonstrates the loss decreasing smoothly and "
            "monotonically toward a minimum as gradient descent proceeds, confirming "
            "correct convergence.",
        param_mod_intro="Effect of varying the learning rate on convergence and final "
            "test accuracy (3000 iterations each):",
        param_mod_headers=["Learning Rate", "Final Loss", "Test Accuracy"],
        param_mod_rows=[[lr, f"{fl:.4f}", f"{a:.4f}"] for lr, fl, a in
                         zip(sweep["learning_rates"], sweep["final_losses"], sweep["accuracies"])],
        param_mod_conclusion="A learning rate that is too small (0.001) has not converged "
            "within 3000 iterations (higher loss, lower accuracy); rates from 0.01 upward "
            "converge to essentially the same optimum, showing the loss surface is well "
            "conditioned once the step size is large enough.",
    )


def q5_cfg():
    m = load_json("q5_pca_tsne")
    sweep = m["perplexity_sweep"]
    return dict(
        exp_no="05", title="PCA and t-SNE — 2D Dimensionality Reduction and Visualization",
        filename="PCA and t-SNE ED408.docx", script_name="scripts/q5_pca_tsne.py",
        aim="Apply PCA and t-SNE to reduce a high-dimensional dataset to 2D and "
            "visualize the resulting class clusters.",
        problem_statement="Project the 10-dimensional standardized pixel feature space "
            "down to 2 dimensions using both a linear technique (PCA) and a non-linear "
            "technique (t-SNE), and visually assess how well edge and non-edge pixels "
            "separate into clusters. Per lab instructions this dataset is used in place "
            "of MNIST. " + DATASET_NOTE,
        algorithm=[
            "Load and standardize the feature dataset.",
            "PCA: compute the covariance matrix of the standardized features.",
            "PCA: compute its eigenvectors/eigenvalues and project data onto the "
            "top 2 principal components (directions of maximum variance).",
            "t-SNE: compute pairwise similarities between points in the 10-D space.",
            "t-SNE: initialize a 2D embedding (via PCA) and iteratively minimize the "
            "KL-divergence between the high-D and low-D similarity distributions.",
            "Plot both 2D projections, coloured by the is_edge label.",
            "Compare cluster separation between PCA and t-SNE.",
        ],
        code='''from sklearn.decomposition import PCA
from sklearn.manifold import TSNE

# PCA - full dataset (linear projection)
pca = PCA(n_components=2, random_state=42)
X_pca = pca.fit_transform(X_scaled)
explained = pca.explained_variance_ratio_

# t-SNE - subsample for tractable runtime (non-linear projection)
tsne = TSNE(n_components=2, perplexity=30,
            random_state=42, init="pca")
X_tsne = tsne.fit_transform(X_scaled[sample_idx])''',
        parameters=[
            ("Feature Set", "10 standardized features"),
            ("PCA sample size", "24,000 (full dataset)"),
            ("t-SNE sample size", m["tsne_sample_size"]),
            ("PCA components", 2),
            ("t-SNE perplexity", m["tsne_perplexity"]),
            ("t-SNE init", "pca"),
        ],
        results=[
            ("PCA var. explained (PC1)", f"{m['pca_explained_variance_ratio'][0]*100:.1f}%"),
            ("PCA var. explained (PC2)", f"{m['pca_explained_variance_ratio'][1]*100:.1f}%"),
            ("PCA total var. explained (2D)", f"{m['pca_total_variance_explained_2d']*100:.1f}%"),
        ],
        graph_file="q5_tsne_2d.png",
        graph_caption="t-SNE 2D projection of the edge-pixel feature space (n=4,000 subsample).",
        x_axis="t-SNE dimension 1", y_axis="t-SNE dimension 2",
        graph_description="The graph demonstrates that t-SNE forms visibly tighter, more "
            "separated edge / non-edge clusters than the PCA projection, since it preserves "
            "local non-linear neighbourhood structure that a linear projection cannot capture.",
        param_mod_intro="Effect of varying t-SNE perplexity on class separation, measured by "
            "the silhouette score of the is_edge label in the resulting 2D embedding "
            "(1,500-point subsample):",
        param_mod_headers=["Perplexity", "Silhouette Score"],
        param_mod_rows=[[p, f"{s:.4f}"] for p, s in
                         zip(sweep["perplexities"], sweep["silhouette_scores"])],
        param_mod_conclusion="Separation peaks around perplexity 15-30 and declines at very "
            "high perplexity (100), because a large perplexity forces t-SNE to consider "
            "distant, less locally-relevant neighbours when building the embedding.",
    )


def q6_cfg():
    m = load_json("q6_outliers")
    sweep = m["threshold_sweep"]
    return dict(
        exp_no="06", title="Outlier Detection and Treatment (Z-score, IQR)",
        filename="Outlier Detection ED408.docx", script_name="scripts/q6_outliers.py",
        aim="Detect and treat outliers in a dataset using the Z-score method and "
            "the IQR method.",
        problem_statement="Identify pixels whose colour/gradient/texture feature values "
            "are statistical outliers using two standard methods, compare how many "
            "outliers each method flags, and treat (cap) them so downstream models are "
            "not skewed by extreme values. " + DATASET_NOTE,
        algorithm=[
            "Load the feature dataset.",
            "Z-score method: for each numeric feature, compute mean and standard "
            "deviation; flag a value as an outlier if |z-score| > 3.",
            "IQR method: for each numeric feature, compute Q1, Q3 and IQR = Q3-Q1; "
            "flag a value as an outlier if it falls outside [Q1-1.5*IQR, Q3+1.5*IQR].",
            "Compare the number of outliers detected by each method, per feature.",
            "Treat outliers by capping (winsorizing) values at the IQR fences.",
            "Visualize feature distributions (box plots) before and after treatment.",
        ],
        code='''def zscore_outliers(s, thresh=3.0):
    z = (s - s.mean()) / s.std()
    return z.abs() > thresh

def iqr_outliers(s, k=1.5):
    q1, q3 = s.quantile(0.25), s.quantile(0.75)
    iqr = q3 - q1
    lower, upper = q1 - k * iqr, q3 + k * iqr
    return (s < lower) | (s > upper), lower, upper

for col in OUTLIER_COLS:
    z_mask = zscore_outliers(df[col])
    iqr_mask, lower, upper = iqr_outliers(df[col])
    # Treatment: winsorize (cap) at the IQR fences
    df_treated[col] = df[col].clip(lower=lower, upper=upper)''',
        parameters=[
            ("Features analyzed", "R, G, B, gray, grad_mag, laplacian, local_std"),
            ("Z-score threshold", 3.0),
            ("IQR multiplier (k)", 1.5),
            ("Treatment method", "Winsorizing (capping at IQR fences)"),
        ],
        results=[
            ("grad_mag outliers (Z-score / IQR)",
             f"{m['per_feature']['grad_mag']['n_outliers_zscore']} / "
             f"{m['per_feature']['grad_mag']['n_outliers_iqr']}"),
            ("laplacian outliers (Z-score / IQR)",
             f"{m['per_feature']['laplacian']['n_outliers_zscore']} / "
             f"{m['per_feature']['laplacian']['n_outliers_iqr']}"),
            ("local_std outliers (Z-score / IQR)",
             f"{m['per_feature']['local_std']['n_outliers_zscore']} / "
             f"{m['per_feature']['local_std']['n_outliers_iqr']}"),
            ("R, G, B, gray outliers", "0 (bounded [0,1] range)"),
        ],
        graph_file="q6_outlier_boxplots.png",
        graph_caption="Box plots of the outlier-prone features before vs. after IQR-capping.",
        x_axis="Feature", y_axis="Feature value",
        graph_description="The graph demonstrates that the long whiskers/extreme points "
            "present before treatment are pulled in to the IQR fences after capping, "
            "without changing the bulk of each distribution.",
        graph_width=5.4,
        param_mod_intro=f"Effect of varying the detection threshold on outlier count for "
            f"'{sweep['demo_column']}' (the feature with the most outliers):",
        param_mod_headers=["Z-threshold", "# Outliers", "IQR multiplier (k)", "# Outliers"],
        param_mod_rows=[[zt, zc, ik, ic] for zt, zc, ik, ic in
                         zip(sweep["z_thresholds"], sweep["z_counts"],
                             sweep["iqr_ks"], sweep["iqr_counts"])],
        param_mod_conclusion="Both methods flag fewer points as the threshold is loosened "
            "(higher Z-threshold or larger IQR multiplier), as expected; IQR consistently "
            "flags more points than Z-score at 'equivalent' settings because grad_mag is "
            "right-skewed, not normally distributed.",
    )


def q7_cfg():
    m = load_json("q7_compare_models")
    lr, svm = m["Logistic Regression"], m["SVM (RBF)"]
    csweep = m["svm_C_sweep"]
    metrics = ["accuracy", "precision", "recall", "f1", "roc_auc"]
    return dict(
        exp_no="07", title="Logistic Regression vs SVM — Metric Comparison",
        filename="Model Comparison ED408.docx", script_name="scripts/q7_compare_models.py",
        aim="Compare Accuracy, Precision, Recall, F1-score and ROC-AUC for Logistic "
            "Regression and SVM on a classification dataset.",
        problem_statement="Train a linear model (Logistic Regression) and a non-linear "
            "model (SVM with RBF kernel) on the same edge-pixel classification task, and "
            "compare their predictive quality across five standard classification "
            "metrics. " + DATASET_NOTE,
        algorithm=[
            "Load and standardize the feature dataset; split into train (80%) / "
            "test (20%).",
            "Train a Logistic Regression classifier on the training set.",
            "Train an SVM (RBF kernel) classifier on the training set.",
            "Generate class predictions and probability/decision scores from both "
            "models on the test set.",
            "Compute Accuracy, Precision, Recall and F1-score for both models.",
            "Compute the ROC curve and AUC for both models.",
            "Compare all five metrics side-by-side.",
        ],
        code='''from sklearn.linear_model import LogisticRegression
from sklearn.svm import SVC
from sklearn.metrics import (accuracy_score, precision_score,
    recall_score, f1_score, roc_auc_score)

log_reg = LogisticRegression(max_iter=2000, random_state=42)
log_reg.fit(X_train, y_train)

svm = SVC(kernel="rbf", C=1.0, gamma="scale",
          probability=True, random_state=42)
svm.fit(X_train, y_train)

for clf in [log_reg, svm]:
    preds = clf.predict(X_test)
    scores = clf.predict_proba(X_test)[:, 1]
    accuracy_score(y_test, preds); precision_score(y_test, preds)
    recall_score(y_test, preds);   f1_score(y_test, preds)
    roc_auc_score(y_test, scores)''',
        parameters=[
            ("Feature Set", "10 standardized features"),
            ("Train / Test Split", "19,200 / 4,800 samples (80/20)"),
            ("Logistic Regression", "max_iter=2000"),
            ("SVM", "kernel=RBF, C=1.0, gamma=scale"),
        ],
        results=[[met.upper(), f"LR={lr[met]:.4f}  |  SVM={svm[met]:.4f}"] for met in metrics],
        graph_file="q7_metric_comparison.png",
        graph_caption="Accuracy, Precision, Recall, F1 and ROC-AUC — Logistic Regression vs SVM.",
        x_axis="Metric", y_axis="Score",
        graph_description="The graph demonstrates SVM (RBF) outperforming Logistic Regression "
            "on every metric, because its non-linear kernel can model the curved decision "
            "boundary between edge and non-edge pixels that a linear model cannot.",
        graph_width=4.8,
        param_mod_intro="Effect of varying the SVM regularization parameter C on Accuracy "
            "and F1-score:",
        param_mod_headers=["C", "Accuracy", "F1-score"],
        param_mod_rows=[[c, f"{a:.4f}", f"{f:.4f}"] for c, a, f in
                         zip(csweep["C_values"], csweep["accuracies"], csweep["f1_scores"])],
        param_mod_conclusion="Both metrics improve as C increases from 0.01 to 10 (a larger "
            "C allows a more complex, less-regularized boundary), then plateau/slightly "
            "decline at C=100, indicating C=10 is close to the sweet spot before overfitting.",
    )


if __name__ == "__main__":
    for cfg_fn in [q1_cfg, q2_cfg, q3_cfg, q4_cfg, q5_cfg, q6_cfg, q7_cfg]:
        build_doc(cfg_fn())
